# -*- coding:utf-8 -*-
# --------------------------------------------------------
# Copyright (C), 2016-2020, lizhe, All rights reserved
# --------------------------------------------------------
# @Name:        convert_rtf.py
# @Author:      lizhe
# @Created:     2023/6/6 - 22:36
# --------------------------------------------------------
import os

from docx import Document

def convert_rtf_to_txt(rtf_path, txt_path):
    # 创建 Word 文档对象
    doc = Document(rtf_path)

    directory, file = os.path.split(rtf_path)

    target_file = os.path.join(directory, file.replace(".rtf", ".docx"))

    doc.save(target_file)

    doc = Document(target_file)

    # 提取纯文本内容
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    # 将文本写入 TXT 文件
    with open(txt_path, 'w', encoding='utf-8') as file:
        file.write(text)


# 设置输入和输出文件夹路径
input_folder = r'C:\Users\lizhe\Desktop\rtf'
output_folder = r'C:\Users\lizhe\Desktop\pic'

# 遍历输入文件夹中的所有文件
for filename in os.listdir(input_folder):
    if filename.endswith('.rtf'):
        # 构建输入和输出文件的完整路径
        rtf_path = os.path.join(input_folder, filename)
        txt_path = os.path.join(output_folder, filename.replace('.rtf', '.txt'))
        try:
            # 转换为纯文本
            convert_rtf_to_txt(rtf_path, txt_path)
        except Exception:
            print(os.path.exists(rtf_path))
            print(f"rtf_path: {rtf_path}")
            print(f"txt_path: {txt_path}")
